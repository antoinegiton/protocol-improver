# text_extractor.py - Extract text from PDF and DOCX files

from pathlib import Path
from typing import Tuple, Dict
import pypdf
from docx import Document

class TextExtractor:
    """Extract text from various document formats"""
    
    def extract(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from PDF file"""
        try:
            text_parts = []
            metadata = {
                'file_type': 'pdf',
                'pages': 0,
                'extraction_method': 'pypdf'
            }
            
            # Read PDF
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                        text_parts.append(page_text)
                
                # Get PDF metadata if available
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from the PDF")
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from DOCX file"""
        try:
            text_parts = []
            metadata = {
                'file_type': 'docx',
                'paragraphs': 0,
                'extraction_method': 'python-docx'
            }
            
            # Read DOCX
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    metadata['paragraphs'] += 1
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            # Get document properties if available
            core_properties = doc.core_properties
            metadata['title'] = core_properties.title or ''
            metadata['author'] = core_properties.author or ''
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from the DOCX")
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def preview_text(self, text: str, max_chars: int = 500) -> str:
        """
        Create a preview of extracted text
        
        Args:
            text: Full text
            max_chars: Maximum characters to show
            
        Returns:
            Preview string
        """
        if len(text) <= max_chars:
            return text
        
        return text[:max_chars] + "..."


# Test function
if __name__ == "__main__":
    # This runs when you execute: python text_extractor.py
    print("Testing TextExtractor...")
    
    extractor = TextExtractor()
    
    # You can test with a sample file
    # test_file = Path("sample_protocol.pdf")
    # text, metadata = extractor.extract(test_file)
    # print(f"Extracted {len(text)} characters")
    # print(f"Metadata: {metadata}")
    
    print("✅ TextExtractor module ready")
